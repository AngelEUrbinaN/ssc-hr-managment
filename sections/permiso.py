from PyQt6.QtWidgets import QWidget, QTableWidgetItem, QMessageBox, QFileDialog
from ui_files.permiso_ui import Ui_Form
import pandas as pd
from docx import Document
import datetime
import locale

class PermisoWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.ui = Ui_Form()
        self.ui.setupUi(self)

        # Set locale for current date
        locale.setlocale(locale.LC_TIME, "Spanish_Mexico.1252")

        # Set ReadOnly the registrar button
        self.ui.registrar_btn.setDisabled(True)

        # Set today date
        current_date = datetime.datetime.now().strftime("%d/%m/%Y")
        self.ui.inicio_input.setDate(datetime.datetime.strptime(current_date, "%d/%m/%Y"))
        self.ui.termino_input.setDate(datetime.datetime.strptime(current_date, "%d/%m/%Y"))
        self.ui.reincorporacion_input.setDate(datetime.datetime.strptime(current_date, "%d/%m/%Y"))

        # Set text dates
        self.ui.inicio_text_input.setText(datetime.datetime.now().strftime("%A %d de %B de %Y"))
        self.ui.termino_text_input.setText(datetime.datetime.now().strftime("%A %d de %B de %Y"))
        self.ui.reincorporacion_text_input.setText(datetime.datetime.now().strftime("%A %d de %B de %Y"))

        # Calculate new date
        self.ui.inicio_input.dateChanged.connect(self.calculate_new_date)
        self.ui.dias_disfrutar_input.valueChanged.connect(self.calculate_new_date)

        # Set text dates on change
        self.ui.inicio_input.dateChanged.connect(self.set_text_date)
        self.ui.termino_input.dateChanged.connect(self.set_text_date)
        self.ui.reincorporacion_input.dateChanged.connect(self.set_text_date)

        # Load recent records when the widget is initialized
        self.load_recent_records()

        # Load employee data if there is a value in the no_emp_input
        self.ui.no_emp_input.returnPressed.connect(self.load_employee_data)

        # Connect registrar button
        self.ui.registrar_btn.clicked.connect(self.create_word_document)

    def load_employee_data(self):
        # Load employee data from Excel
        try:
            df = self.read_main_database()
        except Exception:
            QMessageBox.warning(self, "Leer Base de Datos", "Error al leer la Base de Datos de Empleados.")
            return
        
        no_empleado = self.ui.no_emp_input.text()

        if no_empleado.isdigit():
            no_empleado = int(no_empleado)
            if no_empleado in df['No. DE EMPLEADO'].values:
                try:
                    row = df[df['No. DE EMPLEADO'] == no_empleado].index[0]
                    self.ui.aPaterno_input.setText(df.loc[row, 'APELLIDO PATERNO'])
                    self.ui.aMaterno_input.setText(df.loc[row, 'APELLIDO MATERNO'])
                    self.ui.nombre_input.setText(df.loc[row, 'NOMBRE'])
                    self.ui.area_input.setText(df.loc[row, 'Área Asignada'])
                    self.ui.puesto_input.setText(df.loc[row, 'GRADO'])
                    self.ui.fecha_ingreso_input.setDate(df.loc[row, 'INGRESO'])
                    self.ui.registrar_btn.setDisabled(False)
                except Exception as e:
                    QMessageBox.warning(self, "Error al obtener datos", f'Asegúrate de que el nombre de las columnas sean correctas. Error: {e}')
                    self.ui.registrar_btn.setDisabled(True)
                    self.clear_inputs()
            else:
                QMessageBox.warning(self, "Empleado no encontrado", f'El empleado {no_empleado} no se encuentra en la base de datos.')
                self.ui.registrar_btn.setDisabled(True)
                self.ui.no_emp_input.setText(str(no_empleado))
                self.clear_inputs()

    def get_input_data(self):
        data = {
            "No.Empleado": self.ui.no_emp_input.text(),
            "Nombre": self.ui.nombre_input.text(),
            "Apellido Paterno": self.ui.aPaterno_input.text(),
            "Apellido Materno": self.ui.aMaterno_input.text(),
            "Área": self.ui.area_input.text(),
            "Puesto": self.ui.puesto_input.text(),
            "Fecha Ingreso": self.ui.fecha_ingreso_input.text(),
            "Tipo de Permiso": self.ui.tipo_input.currentText(),
            "Día de Descanso": self.ui.dia_descanso_input.text(),
            "Días a Disfrutar": self.ui.dias_disfrutar_input.value(),
            "Inicio": self.ui.inicio_input.text(),
            "Termino": self.ui.termino_input.text(),
            "Reincorporación": self.ui.reincorporacion_input.text(),
            "Observaciones": self.ui.observ_input.toPlainText()
        }

        # Which type of permiso
        E, CAT, PN, SGDS, P, PNA, PD, PC = self.def_type_of_permiso(data['Tipo de Permiso'])
        data['E'], data['CAT'], data['PN'], data['SGDS'], data['P'], data['PNA'], data['PD'], data['PC'] = E, CAT, PN, SGDS, P, PNA, PD, PC

        # Is sindicalizado
        S, N = self.def_sindicalizado()
        data['S'], data['N'] = S, N

        # Get date for day, month and year
        # Inicio
        day, month, year = self.get_day_month_year(data['Inicio'])
        data['DI'], data['MI'], data['AI'] = day, month, year
        # Termino
        day, month, year = self.get_day_month_year(data['Termino'])
        data['DT'], data['MT'], data['AT'] = day, month, year
        # Reincorporación
        day, month, year = self.get_day_month_year(data['Reincorporación'])
        data['DR'], data['MR'], data['AR'] = day, month, year

        return data
    
    def def_type_of_permiso(self, tipo_permiso):
        E, CAT, PN, SGDS, P, PNA, PD, PC = ' ', ' ', ' ', ' ', ' ', ' ', ' ', ' '

        if tipo_permiso == 'Económico':
            E = 'X'
        elif tipo_permiso == 'Cuidados a terceros':
            CAT = 'X'
        elif tipo_permiso == 'Nupcias':
            PN = 'X'
        elif tipo_permiso == 'Sin goce de sueldo':
            SGDS = 'X'
        elif tipo_permiso == 'Prejubilatoria':
            P = 'X'
        elif tipo_permiso == 'Nacimiento':
            PNA = 'X'
        elif tipo_permiso == 'Defunción':
            PD = 'X'
        elif tipo_permiso == 'Cumpleaños':
            PC = 'X'

        return E, CAT, PN, SGDS, P, PNA, PD, PC
    
    def def_sindicalizado(self):
        S, N = ' ', ' '

        if self.ui.si_input.isChecked():
            S = 'X'
        elif self.ui.no_input.isChecked():
            N = 'X'

        return S, N

    def get_day_month_year(self, date):
        date = datetime.datetime.strptime(date, '%d/%m/%Y')
        day = date.day
        month = date.month
        year = date.year
        return day, month, year
    
    def register_data_on_database(self):
        # Get data from inputs
        data = self.get_input_data()

        # Delete the extra values
        data.pop('E')
        data.pop('CAT')
        data.pop('PN')
        data.pop('SGDS')
        data.pop('P')
        data.pop('PNA')
        data.pop('PD')
        data.pop('PC')
        data.pop('S')
        data.pop('N')

        # Delete the date values
        data.pop('DI')
        data.pop('MI')
        data.pop('AI')
        data.pop('DT')
        data.pop('MT')
        data.pop('AT')
        data.pop('DR')
        data.pop('MR')
        data.pop('AR')

        # Create a new DataFrame with the data
        nuevo = pd.DataFrame([data])
        try:
            df = pd.read_excel("./data/database/permisos.xlsx") if pd.io.common.file_exists("./data/database/permisos.xlsx") else pd.DataFrame()
        except Exception:
            df = pd.DataFrame()

        # Append the new DataFrame to the existing DataFrame
        df = pd.concat([df, nuevo], ignore_index=True)
        df.to_excel("./data/database/permisos.xlsx", index=False)

        self.clear_inputs()
        self.load_recent_records()
    
    def create_word_document(self):
        template = Document('./data/templates/permiso.docx')
        data = self.get_input_data()
        for input in data:
            try:
                data[input] = str(data[input])
                if input == 'Días a Disfrutar':
                    if data[input] == '1':
                        data[input] = data[input] + ' DIA'
                    else:
                        data[input] = data[input] + ' DIAS'
            except Exception:
                pass

        def replace_markers(paragraph, key, value):
            text = "".join(run.text for run in paragraph.runs)
            if f"{{{{{key}}}}}" in text:
                text = text.replace(f"{{{{{key}}}}}", str(value))
                for i, run in enumerate(paragraph.runs):
                    if i == 0:
                        run.text = text
                    else:
                        run.text = ""

        # Reemplazo en párrafos
        for paragraph in template.paragraphs:
            for key, value in data.items():
                replace_markers(paragraph, key, value)

        # Reemplazo en tablas
        for table in template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        for paragraph in cell.paragraphs:
                            replace_markers(paragraph, key, value)

        # Suggested file name
        suggested_name = f"{data['Apellido Paterno'].upper()}_{data['Apellido Materno'].upper()}_{data['Nombre'].upper()}_PERMISO.docx"

        # Ask user where to save
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Document",
            suggested_name,
            "Word Documents (*.docx)"
        )

        if save_path:
            template.save(save_path)
            self.register_data_on_database()
            QMessageBox.information(self, "Saved", f"Document saved at:\n{save_path}")

    def clear_inputs(self):
        self.ui.no_emp_input.clear()
        self.ui.nombre_input.clear()
        self.ui.aPaterno_input.clear()
        self.ui.aMaterno_input.clear()
        self.ui.area_input.clear()
        self.ui.puesto_input.clear()
        self.ui.fecha_ingreso_input.clear()
        self.ui.dia_descanso_input.clear()
        self.ui.dias_disfrutar_input.setValue(0)
        self.ui.inicio_input.clear()
        self.ui.termino_input.clear()
        self.ui.reincorporacion_input.clear()
        self.ui.observ_input.clear()
    
    def read_main_database(self):
        try:
            df = pd.read_excel("./data/database/main.xlsx")
        except Exception as e:
            QMessageBox.warning(self, "Leer Base de Datos", f"Error al leer la Base de Datos de Empleados: {str(e)}")
            return

        return df

    def load_recent_records(self):
        # Load recent records from quinquenio database
        try:
            df = pd.read_excel("./data/database/permisos.xlsx")
        except Exception:
            self.ui.table_display.setRowCount(0)
            self.ui.table_display.setColumnCount(0)
            return

        self.ui.table_display.setRowCount(len(df))
        self.ui.table_display.setColumnCount(len(df.columns))
        self.ui.table_display.setHorizontalHeaderLabels(df.columns)

        # Reverse the DataFrame (most recent records first)
        df = df.iloc[::-1].reset_index(drop=True)

        # Set the table data
        for i in range(len(df)):
            for j in range(len(df.columns)):
                item = QTableWidgetItem(str(df.iat[i, j]))
                self.ui.table_display.setItem(i, j, item)
    
    def format_date_string(self, date_obj):
        # Format string from yyyy-mm-dd hh:mm:ss to dd/mm/yyyy 
        return date_obj.strftime("%d/%m/%Y")
    
    def calculate_new_date(self):
        days = self.ui.dias_disfrutar_input.value() - 1

        termino_date = self.ui.inicio_input.date().toPyDate() + datetime.timedelta(days=days)
        self.ui.termino_input.setDate(termino_date)

        reincorporacion_date = termino_date + datetime.timedelta(days=1)
        self.ui.reincorporacion_input.setDate(reincorporacion_date)

    def set_text_date(self):
        inicio_date = self.ui.inicio_input.date().toPyDate()
        self.ui.inicio_text_input.setText(datetime.datetime.strftime(inicio_date, "%A %d de %B de %Y").capitalize())

        termino_date = self.ui.termino_input.date().toPyDate()
        self.ui.termino_text_input.setText(datetime.datetime.strftime(termino_date, "%A %d de %B de %Y").capitalize())

        reincorporacion_date = self.ui.reincorporacion_input.date().toPyDate()
        self.ui.reincorporacion_text_input.setText(datetime.datetime.strftime(reincorporacion_date, "%A %d de %B de %Y").capitalize())