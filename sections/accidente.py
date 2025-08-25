from PyQt6.QtWidgets import QWidget, QTableWidgetItem, QMessageBox, QFileDialog
from PyQt6.QtCore import QDate, QTime
from ui_files.accidente_ui import Ui_Form
import pandas as pd
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import datetime

class AccidenteWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.ui = Ui_Form()
        self.ui.setupUi(self)

        # Set ReadOnly the registrar button
        self.ui.registrar_btn.setDisabled(True)

        # Load recent records when the widget is initialized
        self.load_recent_records()

        # Set current date
        self.ui.descanso_input.setDate(datetime.datetime.now())
        self.ui.fecha_prt_input.setDate(datetime.datetime.now())
        self.ui.fecha_aviso_input.setDate(datetime.datetime.now())

        # Load employee data if there is a value in the no_emp_input
        self.ui.no_emp_input.returnPressed.connect(self.load_employee_data)

        # Set same date for fecha_prt_input and fecha_aviso_input
        self.ui.fecha_prt_input.dateChanged.connect(self.set_same_date)

        # Connect registrar button
        self.ui.registrar_btn.clicked.connect(self.create_word_document)

    def load_employee_data(self):
        # Load employee data from Excel
        try:
            df = self.read_main_database()
        except Exception:
            QMessageBox.warning(self, "Leer Base de Datos", "Error al leer la Base de Datos de Empleados.")
            return
        
        no_empleado = self.ui.no_emp_input.text()

        if no_empleado.isdigit():
            no_empleado = int(no_empleado)
            if no_empleado in df['No. DE EMPLEADO'].values:
                try:
                    row = df[df['No. DE EMPLEADO'] == no_empleado].index[0]
                    self.ui.nombre_input.setText(df.loc[row, 'NOMBRE COMPLETO'])
                    self.ui.area_input.setText(df.loc[row, 'Adscripción'])
                    self.ui.categoria_input.setText(df.loc[row, 'CLAVE'])
                    self.ui.registrar_btn.setDisabled(False)
                except Exception as e:
                    QMessageBox.warning(self, "Error al obtener datos", f'Asegúrate de que el nombre de las columnas sean correctas. Error: {e}')
                    self.ui.registrar_btn.setDisabled(True)
                    self.clear_inputs()
            else:
                QMessageBox.warning(self, "Empleado no encontrado", f'El empleado {no_empleado} no se encuentra en la base de datos.')
                self.ui.registrar_btn.setDisabled(True)
                self.ui.no_emp_input.setText(str(no_empleado))
                self.clear_inputs()

    def get_input_data(self):
        data = {
            "No.Empleado": self.ui.no_emp_input.text(),
            "Nombre": self.ui.nombre_input.text(),
            "CEGE": self.ui.cege_input.text(),
            "Área": self.ui.area_input.text(),
            "Categoría": self.ui.categoria_input.text(),
            "Horario": self.ui.horario_input_1.text() + " A " + self.ui.horario_input_2.text() + " HRS",
            "Hora del Accidente": self.ui.hora_input.text() + " hrs",
            "Día de Descanso": self.ui.descanso_input.text(),
            "Fecha de PRT": self.ui.fecha_prt_input.text(),
            "Fecha de Aviso": self.ui.fecha_aviso_input.text(),
            "Lugar": self.ui.lugar_input.text(),
            "Comentarios": self.ui.comentarios_input.toPlainText()
        }

        return data
    
    def register_data_on_database(self):
        # Get data from inputs
        data = self.get_input_data()

        # Create a new DataFrame with the data
        nuevo = pd.DataFrame([data])
        try:
            df = pd.read_excel("./data/database/accidentes.xlsx") if pd.io.common.file_exists("./data/database/accidentes.xlsx") else pd.DataFrame()
        except Exception:
            df = pd.DataFrame()

        # Append the new DataFrame to the existing DataFrame
        df = pd.concat([df, nuevo], ignore_index=True)
        df.to_excel("./data/database/accidentes.xlsx", index=False)

        self.clear_inputs()
        self.load_recent_records()

    
    def create_word_document(self):
        template = Document('./data/templates/accidente.docx')
        data = self.get_input_data()
        for input in data:
            try:
                data[input] = str(data[input]) 
            except Exception:
                pass

        def replace_markers(paragraph, key, value):
            text = "".join(run.text for run in paragraph.runs)
            if f"{{{{{key}}}}}" in text:
                text = text.replace(f"{{{{{key}}}}}", str(value))
                for i, run in enumerate(paragraph.runs):
                    if i == 0:
                        run.text = text
                    else:
                        run.text = ""

        # Reemplazo en párrafos
        for paragraph in template.paragraphs:
            for key, value in data.items():
                replace_markers(paragraph, key, value)

        # Reemplazo en tablas
        for table in template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        for paragraph in cell.paragraphs:
                            replace_markers(paragraph, key, value)

        si_selected = self.ui.aceptacion_si_button.isChecked()
        no_selected = self.ui.aceptacion_no_button.isChecked()
        shade_acceptance_cell(template, si_selected, no_selected)

        # Suggested file name
        suggested_name = f"{data['Nombre'].upper()}_ACCIDENTE.docx"

        # Ask user where to save
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Document",
            suggested_name,
            "Word Documents (*.docx)"
        )

        if save_path:
            template.save(save_path)
            self.register_data_on_database()
            QMessageBox.information(self, "Saved", f"Document saved at:\n{save_path}")

    def clear_inputs(self):
        self.ui.no_emp_input.clear()
        self.ui.nombre_input.clear()
        self.ui.lugar_input.clear()
        self.ui.area_input.clear()
        self.ui.categoria_input.clear()
        self.ui.horario_input_1.setTime(QTime(0, 0))
        self.ui.horario_input_2.setTime(QTime(0, 0))
        self.ui.hora_input.setTime(QTime(0, 0))
        self.ui.comentarios_input.clear()
        self.ui.comentarios_label.setText("Se anexa formato de calificación de probable accidente de trabajo ST- 7 y tarjeta informativa")
        self.ui.aceptacion_si_button.setChecked(False)
        self.ui.aceptacion_no_button.setChecked(False)
        today_qdate = QDate.currentDate()
        self.ui.descanso_input.setDate(today_qdate)
        self.ui.fecha_prt_input.setDate(today_qdate)
        self.ui.fecha_aviso_input.setDate(today_qdate)
    
    def read_main_database(self):
        try:
            df = pd.read_excel("./data/database/main.xlsx")
        except Exception as e:
            QMessageBox.warning(self, "Leer Base de Datos", f"Error al leer la Base de Datos de Empleados: {str(e)}")
            return

        return df

    def load_recent_records(self):
        # Load recent records from quinquenio database
        try:
            df = pd.read_excel("./data/database/accidentes.xlsx")
        except Exception:
            self.ui.table_display.setRowCount(0)
            self.ui.table_display.setColumnCount(0)
            return

        self.ui.table_display.setRowCount(len(df))
        self.ui.table_display.setColumnCount(len(df.columns))
        self.ui.table_display.setHorizontalHeaderLabels(df.columns)

        # Reverse the DataFrame (most recent records first)
        df = df.iloc[::-1].reset_index(drop=True)

        # Set the table data
        for i in range(len(df)):
            for j in range(len(df.columns)):
                item = QTableWidgetItem(str(df.iat[i, j]))
                self.ui.table_display.setItem(i, j, item)

    def set_same_date(self):
        self.ui.fecha_aviso_input.setDate(self.ui.fecha_prt_input.date())

def shade_cell(cell, hex_color="D9D9D9"):
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
        )

def clear_cell_shading(cell):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="FFFFFF"/>')
    )

def find_cells_by_text(doc: Document, targets: set[str]):
    """Return list of (table_idx, row_idx, cell_idx, cell) for any cell whose text matches one of targets."""
    hits = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if cell.text.strip().upper() in targets:
                    hits.append((ti, ri, ci, cell))
    return hits
    
def shade_acceptance_cell(doc: Document, si_selected: bool, no_selected: bool):
    header_key = "ACEPTACIÓN DEL P. R.T"          # exact as appears in the template
    targets     = {"SÍ", "SI", "NO"}              # allow both with/without accent
    target_row = None
    target_table = None
    for table in doc.tables:
        for row in table.rows:
            # Find the row that contains the header cell
            if any(header_key in cell.text for cell in row.cells):
                target_row = row
                target_table = table
                break
        if target_row:
            break
    if not target_row:
        return  # header row not found; silently skip
    # Inside that row, locate the “SÍ/NO” cells
    si_cell = no_cell = None
    for cell in target_row.cells:
        t = cell.text.strip().upper()
        if t in {"SÍ", "SI"}:
            si_cell = cell
        elif t == "NO":
            no_cell = cell
    # Clear both first (so the file is idempotent)
    if si_cell: clear_cell_shading(si_cell)
    if no_cell: clear_cell_shading(no_cell)
    # Shade the one selected
    if si_selected and si_cell:
        shade_cell(si_cell)
    elif no_selected and no_cell:
        shade_cell(no_cell)