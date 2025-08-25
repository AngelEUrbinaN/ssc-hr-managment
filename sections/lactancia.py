from PyQt6.QtWidgets import QWidget, QTableWidgetItem, QMessageBox, QFileDialog
from ui_files.lactancia_ui import Ui_Form
import pandas as pd
import locale
import datetime
from dateutil.relativedelta import relativedelta
from docx import Document

class LactanciaWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.ui = Ui_Form()
        self.ui.setupUi(self)

        # Set ReadOnly the not-editable inputs
        self.ui.nombre_input.setReadOnly(True)
        self.ui.aPaterno_input.setReadOnly(True)
        self.ui.aMaterno_input.setReadOnly(True)
        self.ui.area_input.setReadOnly(True)
        self.ui.puesto_input.setReadOnly(True)

        # Set ReadOnly the registrar button
        self.ui.registrar_btn.setDisabled(True)

        # Load recent records when the widget is initialized
        self.load_recent_records()

        # Load employee data if there is a value in the no_emp_input
        self.ui.no_emp_input.returnPressed.connect(self.load_employee_data)

        # Set current date
        self.set_current_date()
        self.calculate_maternity_leave()

        # Calculate maternity leave
        self.ui.mat_inicio_input.dateChanged.connect(self.calculate_maternity_leave)

        # Connect registrar button
        self.ui.registrar_btn.clicked.connect(self.create_word_document)

        # Set locale for current date
        locale.setlocale(locale.LC_TIME, 'es_MX.UTF-8')

    def set_current_date(self):
        # Set current date
        current_date = datetime.datetime.now().strftime("%d/%m/%Y")
        self.ui.mat_inicio_input.setDate(datetime.datetime.strptime(current_date, "%d/%m/%Y"))

    def load_employee_data(self):
        # Load employee data from Excel
        try:
            df = self.read_main_database()
        except Exception:
            QMessageBox.warning(self, "Leer Base de Datos", "Error al leer la Base de Datos de Empleados.")
            return
        
        no_empleado = self.ui.no_emp_input.text()

        if no_empleado.isdigit():
            no_empleado = int(no_empleado)
            if no_empleado in df['No. DE EMPLEADO'].values:
                row = df[df['No. DE EMPLEADO'] == no_empleado].index[0]
                self.ui.aPaterno_input.setText(df.loc[row, 'APELLIDO PATERNO'])
                self.ui.aMaterno_input.setText(df.loc[row, 'APELLIDO MATERNO'])
                self.ui.nombre_input.setText(df.loc[row, 'NOMBRE'])
                self.ui.area_input.setText(df.loc[row, 'Área Asignada'])
                self.ui.puesto_input.setText(df.loc[row, 'GRADO'])
                self.ui.registrar_btn.setDisabled(False)
            else:
                QMessageBox.warning(self, "Empleado no encontrado", f'El empleado {no_empleado} no se encuentra en la base de datos.')
                self.ui.registrar_btn.setDisabled(True)
                self.ui.no_emp_input.setText(str(no_empleado))
                self.clear_inputs()
        
    def calculate_maternity_leave(self):
        # Get QDate and convert to datetime.date
        qdate_start = self.ui.mat_inicio_input.date()
        maternity_start_date = qdate_start.toPyDate()

        # Maternity leave: 84 days (12 weeks)
        maternity_end_date = maternity_start_date + datetime.timedelta(days=83)
        lactation_start_date = maternity_end_date + datetime.timedelta(days=1)

        # Lactation leave: +6 calendar months
        lactation_end_date = lactation_start_date + relativedelta(months=6)

        # Set formatted dates in the UI
        self.ui.mat_termino_input.setDate(maternity_end_date)
        self.ui.lac_inicio_input.setDate(lactation_start_date)
        self.ui.lac_termino_input.setDate(lactation_end_date)

    def get_input_data(self):
        data = {
            "Fecha": datetime.datetime.now().strftime("%d de %B de %Y"),
            "No.Empleado": self.ui.no_emp_input.text(),
            "Nombre": self.ui.nombre_input.text(),
            "Apellido Paterno": self.ui.aPaterno_input.text(),
            "Apellido Materno": self.ui.aMaterno_input.text(),
            "Área": self.ui.area_input.text(),
            "Puesto": self.ui.puesto_input.text(),
            "Inicio M.": self.ui.mat_inicio_input.text(),
            "Fin M.": self.ui.mat_termino_input.text(),
            "Inicio L.": self.ui.lac_inicio_input.text(),
            "Fin L.": self.ui.lac_termino_input.text(),
            "Horas L.": self.ui.hor_lac_input.text(),
            "Observaciones": self.ui.observ_input.toPlainText()
        }

        return data
    
    def register_data_on_database(self):
        # Get data from inputs
        data = self.get_input_data()

        # Create a new DataFrame with the data
        nuevo = pd.DataFrame([data])
        try:
            df = pd.read_excel("./data/database/lactancias.xlsx") if pd.io.common.file_exists("./data/database/lactancias.xlsx") else pd.DataFrame()
        except Exception:
            df = pd.DataFrame()

        # Append the new DataFrame to the existing DataFrame
        df = pd.concat([df, nuevo], ignore_index=True)
        df.to_excel("./data/database/lactancias.xlsx", index=False)

        self.clear_inputs()
        self.load_recent_records()
    
    def create_word_document(self):
        template = Document('./data/templates/lactancia.docx')
        data = self.get_input_data()
        for input in data:
            try:
                data[input] = str(data[input])
                data[input] = data[input].upper()
            except Exception:
                pass

        def replace_markers(paragraph, key, value):
            text = "".join(run.text for run in paragraph.runs)
            if f"{{{{{key}}}}}" in text:
                text = text.replace(f"{{{{{key}}}}}", str(value))
                for i, run in enumerate(paragraph.runs):
                    if i == 0:
                        run.text = text
                    else:
                        run.text = ""

        # Reemplazo en párrafos
        for paragraph in template.paragraphs:
            for key, value in data.items():
                replace_markers(paragraph, key, value)

        # Reemplazo en tablas
        for table in template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        for paragraph in cell.paragraphs:
                            replace_markers(paragraph, key, value)

        # Suggested file name
        suggested_name = f"{data['Apellido Paterno']}_{data['Apellido Materno']}_{data['Nombre']}_LACTANCIA.docx"

        # Ask user where to save
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Document",
            suggested_name,
            "Word Documents (*.docx)"
        )

        if save_path:
            template.save(save_path)
            self.register_data_on_database()
            QMessageBox.information(self, "Saved", f"Document saved at:\n{save_path}")

    def clear_inputs(self):
        self.ui.aPaterno_input.clear()
        self.ui.aMaterno_input.clear()
        self.ui.nombre_input.clear()
        self.ui.area_input.clear()
        self.ui.no_emp_input.clear()
        self.ui.puesto_input.clear()
        self.ui.hor_lac_input.clear()
        self.ui.observ_input.clear()
        self.set_current_date()
        self.calculate_maternity_leave()
    
    def read_main_database(self):
        try:
            df = pd.read_excel("./data/database/main.xlsx")
        except Exception as e:
            QMessageBox.warning(self, "Leer Base de Datos", f"Error al leer la Base de Datos de Empleados: {str(e)}")
            return

        return df

    def load_recent_records(self):
        # Load recent records from quinquenio database
        try:
            df = pd.read_excel("./data/database/lactancias.xlsx")
        except Exception:
            self.ui.table_display.setRowCount(0)
            self.ui.table_display.setColumnCount(0)
            return

        self.ui.table_display.setRowCount(len(df))
        self.ui.table_display.setColumnCount(len(df.columns))
        self.ui.table_display.setHorizontalHeaderLabels(df.columns)

        # Reverse the DataFrame (most recent records first)
        df = df.iloc[::-1].reset_index(drop=True)

        # Set the table data
        for i in range(len(df)):
            for j in range(len(df.columns)):
                item = QTableWidgetItem(str(df.iat[i, j]))
                self.ui.table_display.setItem(i, j, item)